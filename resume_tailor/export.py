import io
import re
from pathlib import Path


def markdown_to_docx(markdown_text: str) -> bytes:
    """Convert markdown text to a DOCX file. Returns bytes."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = markdown_text.split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            p = doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_formatted_text(paragraph, text: str):
    """Parse simple markdown bold (**text**) and add runs to a paragraph."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _sanitize_for_pdf(text: str) -> str:
    """Replace common non-Latin-1 characters for PDF compatibility."""
    replacements = {
        '\u2014': '--',   # em dash
        '\u2013': '-',    # en dash
        '\u2018': "'",    # left single quote
        '\u2019': "'",    # right single quote
        '\u201c': '"',    # left double quote
        '\u201d': '"',    # right double quote
        '\u2026': '...',  # ellipsis
        '\u2022': '-',    # bullet
        '\u00b7': '-',    # middle dot
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text.encode('latin-1', 'replace').decode('latin-1')


def markdown_to_pdf(markdown_text: str) -> bytes:
    """Convert markdown text to a PDF file. Returns bytes."""
    from fpdf import FPDF

    markdown_text = _sanitize_for_pdf(markdown_text)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)

    lines = markdown_text.split("\n")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            pdf.ln(4)
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 18)
            pdf.cell(0, 10, stripped[2:].strip(), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=11)
        elif stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 8, stripped[3:].strip(), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=11)
        elif stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 7, stripped[4:].strip(), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=11)
        elif stripped == "---":
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            clean_text = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped[2:].strip())
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 6, f"    - {clean_text}", new_x="LMARGIN", new_y="NEXT")
        else:
            clean_text = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 6, clean_text, new_x="LMARGIN", new_y="NEXT")

    return pdf.output()
